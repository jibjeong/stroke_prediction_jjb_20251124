"""
Convert Markdown to DOCX
Converts 03_results.md to 03_results.docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a new run object
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color (blue) and underline
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def parse_table(lines):
    """Parse markdown table and return rows"""
    rows = []
    for line in lines:
        # Skip separator lines
        if '|---' in line or '|====' in line:
            continue
        # Parse table row
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    """Add a table to the document"""
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            # Clean markdown formatting
            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)  # Bold
            cell_text = re.sub(r'\*(.*?)\*', r'\1', cell_text)      # Italic
            cell.text = cell_text

            # Make header row bold
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def convert_md_to_docx(md_file, output_file):
    """Convert markdown file to docx"""
    # Create document
    doc = Document()

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    table_lines = []
    in_table = False
    in_code_block = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines at start
        if not line and i == 0:
            i += 1
            continue

        # Code block
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            if not doc.paragraphs or doc.paragraphs[-1].text:
                p = doc.add_paragraph()
                p.style = 'Intense Quote'
            doc.paragraphs[-1].text += line + '\n'
            i += 1
            continue

        # Detect table
        if '|' in line and not in_table:
            in_table = True
            table_lines = []

        if in_table:
            if '|' in line:
                table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                in_table = False
                rows = parse_table(table_lines)
                add_table_to_doc(doc, rows)
                doc.add_paragraph()  # Add space after table
                table_lines = []

        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)

        # Horizontal rule
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 80)

        # Images
        elif line.startswith('!['):
            # Parse: ![alt text](path)
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)

                # Add figure caption
                p = doc.add_paragraph()
                p.add_run(alt_text).bold = True

                # Try to add image (skip if not found)
                full_path = Path(md_file).parent / image_path
                if full_path.exists():
                    try:
                        doc.add_picture(str(full_path), width=Inches(6))
                    except:
                        p = doc.add_paragraph()
                        p.add_run(f'[Image: {alt_text}]').italic = True
                else:
                    p = doc.add_paragraph()
                    p.add_run(f'[Image: {alt_text} - Not Found]').italic = True

        # Lists
        elif line.startswith('- ') or re.match(r'^\d+\.', line):
            text = re.sub(r'^[-\d]+\.\s*', '', line)
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')

        # Regular paragraph
        elif line:
            # Skip lines that are part of special formatting
            if line.startswith('**') and line.endswith('**:'):
                # Section label
                text = line.strip('*').strip(':')
                p = doc.add_paragraph()
                p.add_run(text).bold = True
            else:
                # Regular text with inline formatting
                p = doc.add_paragraph()

                # Split by bold/italic markers
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part.strip('*')).bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        p.add_run(part.strip('*')).italic = True
                    else:
                        p.add_run(part)

        i += 1

    # Save document
    doc.save(output_file)
    print(f"✓ Converted {md_file} to {output_file}")


if __name__ == "__main__":
    # Input and output paths
    md_file = Path("paper/sections/03_results.md")
    output_file = Path("paper/sections/03_results.docx")

    # Convert
    convert_md_to_docx(md_file, output_file)

    print(f"\n✓ Conversion completed successfully!")
    print(f"  Input:  {md_file}")
    print(f"  Output: {output_file}")
